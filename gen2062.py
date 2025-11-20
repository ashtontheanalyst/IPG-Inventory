from pathlib import Path
from docx import Document
from copy import deepcopy
import pandas as pd
from math import floor
import time


# Directory help/paths
BASE_DIR = Path(__file__).resolve().parent
DATA_CSV = BASE_DIR / "data" / "Inventory.csv"
FIRST_2062 = BASE_DIR / "docs" / "2062p1.docx"
CONT_2062 = BASE_DIR / "docs" / "2062cont.docx"
OUTPUT_DIR = BASE_DIR / "docs" / "filled"




# Helper function to generate the DOCx
def generate2062docx():
    # DOC is the overall doc that will get sent to the front end, for now we append the first p of 2062
    DOC = Document(FIRST_2062)
    first2062 = DOC.tables[0]


    # Defaults/States
    f2062startRowInd = 4
    f2062RowMax = 23
    rowF = f2062startRowInd
    rowC = 0
    CONT_DOC = None              # This is where a cont doc will be held if items exceed the amount of first page
    cont2062 = None              # This is the table for the cont doc

    
    # Read the CSV into a df, pull out the items name and serial and put it in the items df if it's been scanned
    df = pd.read_csv(DATA_CSV, dtype=str)
    items =(
        df[df["inventoried"] == "T"][["category_name", "serial_number"]]
        .fillna("")         # Any NaN values we make an empty string
        .astype(str)
        .sort_values("category_name")
    )


    # Editing the header of page 1 2062
    first2062.cell(0, 5).text = "FROM:\nBCDC - IPG"
    first2062.cell(0, 12).text = f"TO:"
    first2062.cell(1, 10).text = f"DATE:\n{time.localtime().tm_mon}/{time.localtime().tm_mday}/{time.localtime().tm_year}"
    first2062.cell(1, 15).text = f"QUANTITY:\n{len(items)}"

    
    # Amount of pages needed, first one holds 16, cont pages hold 19 items
    if len(items) <= 19:
        amountPages = 1
    else:
        contItems = len(items) - 19         # Get rid of first page items for math
        amountPages = floor(contItems / 23) + 1
        print(amountPages)

    
    # This is the name from i-1, this is compared to the current i value to see if we can + the quant or make a new row
    nameMin1 = ""

    
    # Editing doc(s)
    for i in range(len(items)):
        # Values, pulled from the items df that was narrowed down
        name = items.iloc[i]["category_name"]
        serial = items.iloc[i]["serial_number"]

        # If the name of the current item matches the last item
        if name == nameMin1:
            # Ensure at least one row exists
            if rowF < f2062RowMax:
                # First Page
                if rowF <= f2062RowMax:
                    rowFmin1 = rowF - 1
                    serialCell = first2062.cell(rowFmin1, 2)
                    qtyCell = first2062.cell(rowFmin1, 10)
                    
                    # Append the serial #
                    if serialCell.text.strip():
                        serialCell.text = serialCell.text + f", {serial}"
                    else:
                        serialCell.text = serial

                    # Inc quantity
                    currQty = qtyCell.text.strip()
                    currQty = int(currQty) if currQty else 0
                    qtyCell.text = str(currQty + 1)

                    # IMPORTANT, this moves us to the next instead of making a new row
                    continue
            
            # Continuation page(s), same logic as above
            if CONT_DOC is not None and rowC > 0:
                rowCmin1 = rowC - 1
                serialCell = cont2062.cell(rowCmin1, 2)
                qtyCell = cont2062.cell(rowCmin1, 10)

                if serialCell.text.strip():
                    serialCell.text = serialCell.text + f", {serial}"
                else:
                    serialCell.text = serial

                currQty = qtyCell.text.strip()
                currQty = int(currQty) if currQty else 0
                qtyCell.text = str(currQty + 1)

                continue
        
        # Inc name
        nameMin1 = name


        # This is for if we don't have a match, make a new row
        # Page 1
        if rowF < f2062RowMax:
            # Edit the specific cell(row, col) text with narrowed down df values
            first2062.cell(rowF, 0).text = name
            first2062.cell(rowF, 2).text = serial
            first2062.cell(rowF, 10).text = "1"
            rowF += 1

        # Continuation page(s)
        else:
            # Make the cont doc and fill it to the capacity, once filled then append it to main and clear
            # for the next iteration if necessary
            if CONT_DOC is None:
                CONT_DOC = Document(CONT_2062)
                cont2062 = CONT_DOC.tables[0]

            cont2062.cell(rowC, 0).text = name
            cont2062.cell(rowC, 2).text = serial
            cont2062.cell(rowC, 10).text = "1"
            rowC += 1

            # Once the continuation page is full, append to the main doc and reset values
            if rowC == 23:
                if CONT_DOC is not None:
                    for element in CONT_DOC.element.body:
                        DOC.element.body.append(deepcopy(element))

                rowC = 0
                CONT_DOC = None
                cont2062 = None
    

    # Save file
    output_path = OUTPUT_DIR / f"Generated2062.docx"
    DOC.save(output_path)
    return output_path




# Helper function to inspect the doc, basically gives you an x-ray of all the elements on a docx file
def inspectDoc(path):
    doc = Document(path)

    print("<===== DOC Structure =====>")

    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx} ---")
        print(f"Rows: {len(table.rows)}")
        print(f"Cols: {len(table.columns)}")

        for r_idx, row in enumerate(table.rows):
            cell_texts = [cell.text.strip() for cell in row.cells]
            print(f" Row {r_idx}: {cell_texts}")




''' 2062p1.docx x-ray:
<===== DOC Structure =====>

--- TABLE 0 ---
Rows: 21
Cols: 17
 Row 0: ['HAND RECEIPT/ANNEX NUMBER\nFor use of this form, se DA PAM 710-2-1.\nThe Proponent agency is ODCSLOG.', 'HAND RECEIPT/ANNEX NUMBER\nFor use of this form, se DA PAM 710-2-1.\nThe Proponent agency is ODCSLOG.', 'HAND RECEIPT/ANNEX NUMBER\nFor use of this form, se DA PAM 710-2-1.\nThe Proponent agency is ODCSLOG.', 'HAND RECEIPT/ANNEX NUMBER\nFor use of this form, se DA PAM 710-2-1.\nThe Proponent agency is ODCSLOG.', 'FROM:', 'FROM:', 'TO:', 'TO:', 'TO:', 'TO:', 'TO:', 'TO:', 'TO:', 'TO:', 'HAND RECEIPT NUMBER', 'HAND RECEIPT NUMBER', 'HAND RECEIPT NUMBER']
 Row 1: ['FOR ANNEX/CR ONLY', 'END ITEM STOCK NUMBER', 'END ITEM STOCK NUMBER', 'END ITEM DESCRIPTION', 'END ITEM DESCRIPTION', 'PUBLICATION NUMBER', 'PUBLICATION NUMBER', 'PUBLICATION NUMBER', 'PUBLICATION NUMBER', 'PUBLICATION NUMBER', 'PUBLICATION DATE', 'PUBLICATION DATE', 'PUBLICATION DATE', 'PUBLICATION DATE', 'QUANTITY', 'QUANTITY', 'QUANTITY']
 Row 2: ['STOCKNUMBER\n\na.', 'STOCKNUMBER\n\na.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', '*\nc.', 'SEC\nd.', 'UI\ne.', 'QTY AUTH\nf.', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY']
 Row 3: ['STOCKNUMBER\n\na.', 'STOCKNUMBER\n\na.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', 'ITEM DESCRIPTION\n\nb.', '*\nc.', 'SEC\nd.', 'UI\ne.', 'QTY AUTH\nf.', 'A', 'B', 'C', 'D', 'E', 'F']
 Row 4: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 5: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 6: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 7: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 8: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 9: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 10: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 11: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 12: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 13: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 14: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 15: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 16: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 17: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 18: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 19: ['', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
 Row 20: ['* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES', '* WHEN USED AS A:\nHAND RECEIPT, enter Hand Receipt Annex number\nHAND RECEIPT FOR QUARTERS FURNITURE, enter Condition Codes\nHAND RECEIPT ANNEX/COMPONENTS RECEIPT, enter Accounting Requirements Code (ARC).\nPAGE 1 OF 1 PAGES']
'''


''' 2062cont.docx x-ray:
<===== DOC Structure =====>

--- TABLE 0 ---
Rows: 21
Cols: 12
 Row 0: ['STOCKNUMBER\n\na.', 'ITEM DESCRIPTION\n\nb.', '*\nc.', 'SEC\nd.', 'UI\ne.', 'QTY AUTH\nf.', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY', 'g. \tQUANTITY']
 Row 1: ['STOCKNUMBER\n\na.', 'ITEM DESCRIPTION\n\nb.', '*\nc.', 'SEC\nd.', 'UI\ne.', 'QTY AUTH\nf.', 'A', 'B', 'C', 'D', 'E', 'F']
 Row 2: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 3: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 4: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 5: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 6: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 7: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 8: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 9: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 10: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 11: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 12: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 13: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 14: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 15: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 16: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 17: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 18: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 19: ['', '', '', '', '', '', '', '', '', '', '', '']
 Row 20: ['', '', '', '', '', '', '', '', '', '', '', '']
'''